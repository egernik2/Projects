import docx # Импорт библиотеки для работы с файлами формата Docx

TEMPLATE_NAME = 'NIR.docx' # Имя файла шаблона
FONT_NAME = 'Times New Roman' # Имя шрифта
FONT_SIZE = 14 # Размер шрифта
FIO_FILE_NAME = 'fio.list'

def dativ_male(word):
    endings = {
        'а': 'е',
        'я': 'е',
        'о': 'о',
        'е': 'у',
        'ь': 'ю'
    }
    if word[-2:] == 'ий':
        return word[:-2] + 'ию'
    if word[-2:] == 'ай':
        return word[:-2] + 'аю'
    if word[-1] in endings:
        return word[:-1] + endings[word[-1]]
    return word + 'у'

def dativ_female(word, is_second_name):
    endings1 = {
        'а': 'е',
        'я': 'и',
        'о': 'у',
        'е': 'у',
        'ь': 'ю'
    }

    endings2 = {
        'а': 'ой',
        'я': 'е',
        'о': 'е',
        'е': 'у',
        'ь': 'ю'
    }
    if is_second_name:
        if word[-1] in endings2:
            return word[:-1] + endings2[word[-1]]
    else:
        if word[-1] in endings1:
            return word[:-1] + endings1[word[-1]]
    return word + 'у'


with open(FIO_FILE_NAME, 'r') as f: # Открываем файл на чтение
    fio_list = f.readlines() # Считываем файл

letters = ['е', 'а', 'о', 'э', 'и', 'ю', 'я', 'ы', 'у']
iskl = ['Никита', 'Илья', 'Данила', 'Кузьма']

for item in fio_list: # Пробегаемся по словарю...
    doc = docx.Document(docx = TEMPLATE_NAME) # Открываем файл на чтение
    style = doc.styles['Normal'] # Выставляем стиль текста на нормальный
    font = style.font # Объявляем экземпляр шрифта
    font.name = FONT_NAME # Имя шрифта
    font.size = docx.shared.Pt(FONT_SIZE) # Размер шрифта
    second_name, first_name, middle_name = item.split(' ') # Разбиваем строку на фамилию, имя и отчество
    middle_name = middle_name.strip()
    if first_name[-1] in letters and first_name not in iskl:
        second_name_dat = dativ_female(second_name, True) # Склоняем фамилию
        first_name_dat = dativ_female(first_name, False) # Склоняем имя
        middle_name_dat = dativ_female(middle_name, False) # Склоняем отчество
    else:
        second_name_dat = dativ_male(second_name) # Склоняем фамилию
        first_name_dat = dativ_male(first_name) # Склоняем имя
        middle_name_dat = dativ_male(middle_name) # Склоняем отчество
    fio_dat = '{} {} {}'.format(second_name_dat, first_name_dat, middle_name_dat)

    for p in doc.paragraphs: # Для каждого параграфа... Параграфы вне таблицы
        if p.text.find('%ZAMENA2%') >= 0: # Если нашли в тексте параграфа совпадение...
            p.text = p.text.replace('%ZAMENA2%', '{} {}'.format(first_name, middle_name)) # Выполняем замену

    for table in doc.tables: # перебираем все таблицы в документе
        for row in table.rows: # перебираем все строки в таблице
            for cell in row.cells: # перебираем все ячейки в строке
                cell_content = cell.text # получаем содержимое ячейки
                if '%ZAMENA1%' in cell_content: # проверяем, содержит ли ячейка искомое значение
                    style = cell.paragraphs.runs[0].style
                    new_content = cell_content.replace('%ZAMENA1%', fio_dat, style=style) # заменяем искомое значение на новое
                    cell.text = new_content # устанавливаем новое значение в ячейку



    doc.save('{}.docx'.format(fio_dat)) # Сохраняем файл
print ('Done')